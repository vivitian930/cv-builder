import argparse
import json
from docx import Document
from docx.shared import Pt, Inches
from typing import List, Optional
from pydantic import BaseModel

class PersonalInfo(BaseModel):
    name: Optional[str]
    email: Optional[str]
    phone: Optional[str]

class Education(BaseModel):
    degree: Optional[str]
    university: Optional[str]
    graduation_date: Optional[str]

class WorkExperience(BaseModel):
    company: Optional[str]
    position: Optional[str]
    start_date: Optional[str]
    end_date: Optional[str]
    summary: Optional[str]
    bullet_points: Optional[List[str]]

class ResearchExperience(BaseModel):
    organization: Optional[str]
    position: Optional[str]
    start_date: Optional[str]
    end_date: Optional[str]
    summary: Optional[str]
    bullet_points: Optional[List[str]]

class ProjectExperience(BaseModel):
    name: Optional[str]
    start_date: Optional[str]
    end_date: Optional[str]
    summary: Optional[str]
    bullet_points: Optional[List[str]]

class UserInterface:
    def __init__(self):
        self.personal_info = PersonalInfo(name=None, email=None, phone=None)
        self.education = []
        self.work_experience = []
        self.research_experience = []
        self.project_experience = []
        self.personal_statement = ""

    def capture_personal_info(self):
        print("Please provide your personal information:")
        self.personal_info.name = input("Name: ")
        self.personal_info.email = input("Email: ")
        self.personal_info.phone = input("Phone: ")

    def capture_education(self):
        print("Please provide your education details (press enter to skip):")
        while True:
            degree = input("Degree: ")
            if not degree:
                break

            university = input("University: ")
            graduation_date = input("Graduation Date (MM/YYYY): ")

            self.education.append(Education(degree=degree, university=university, graduation_date=graduation_date))

    def capture_work_experience(self):
        print("Please provide your work experience (press enter to skip):")
        while True:
            company = input("Company: ")
            if not company:
                break

            position = input("Position: ")
            start_date = input("Start Date (MM/YYYY): ")
            end_date = input("End Date (MM/YYYY): ")
            summary = input("Summary: ")

            bullet_points = []
            while True:
                bullet_point = input("Bullet Point (press enter to skip): ")
                if not bullet_point:
                    break
                bullet_points.append(bullet_point)

            self.work_experience.append(WorkExperience(
                company=company,
                position=position,
                start_date=start_date,
                end_date=end_date,
                summary=summary,
                bullet_points=bullet_points
            ))

    def capture_research_experience(self):
        print("Please provide your research experience (press enter to skip):")
        while True:
            organization = input("Organization: ")
            if not organization:
                break

            position = input("Position: ")
            start_date = input("Start Date (MM/YYYY): ")
            end_date = input("End Date (MM/YYYY): ")
            summary = input("Summary: ")

            bullet_points = []
            while True:
                bullet_point = input("Bullet Point (press enter to skip): ")
                if not bullet_point:
                    break
                bullet_points.append(bullet_point)

            self.research_experience.append(ResearchExperience(
                organization=organization,
                position=position,
                start_date=start_date,
                end_date=end_date,
                summary=summary,
                bullet_points=bullet_points
            ))

    def capture_project_experience(self):
        print("Please provide your project experience (press enter to skip):")
        while True:
            name = input("Project Name: ")
            if not name:
                break

            start_date = input("Start Date (MM/YYYY): ")
            end_date = input("End Date (MM/YYYY): ")
            summary = input("Summary: ")

            bullet_points = []
            while True:
                bullet_point = input("Bullet Point (press enter to skip): ")
                if not bullet_point:
                    break
                bullet_points.append(bullet_point)

            self.project_experience.append(ProjectExperience(
                name=name,
                start_date=start_date,
                end_date=end_date,
                summary=summary,
                bullet_points=bullet_points
            ))

    def capture_personal_statement(self):
        print("Please provide your personal statement (press enter to skip):")
        self.personal_statement = input("Personal Statement: ")

    def display_resume(self):
        doc = Document()

        # Set document font size and margin
        doc.styles['Normal'].font.size = Pt(10)
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.25)  # Adjust top margin to make space smaller
            section.bottom_margin = Inches(0.25)  # Adjust bottom margin to make space smaller

        # Add personal information
        doc.add_heading("Personal Information", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Name: {self.personal_info.name or ''}").bold = True
        p.add_run(" | ")
        p.add_run(f"Email: {self.personal_info.email or ''}").bold = True
        p.add_run(" | ")
        p.add_run(f"Phone: {self.personal_info.phone or ''}").bold = True

        # Add education details
        doc.add_heading("Education", level=1)
        for education in self.education:
            doc.add_heading(f"{education.degree or ''}", level=2)
            doc.add_paragraph(f"University: {education.university or ''}")
            doc.add_paragraph(f"Graduation Date: {education.graduation_date or ''}")

        # Add work experience
        doc.add_heading("Work Experience", level=1)
        for experience in self.work_experience:
            doc.add_heading(f"{experience.position or ''} at {experience.company or ''}", level=2)
            doc.add_paragraph(f"{experience.start_date or ''} - {experience.end_date or ''}")
            doc.add_paragraph(f"Summary: {experience.summary or ''}")
            for bullet_point in experience.bullet_points or []:
                doc.add_paragraph(f"• {bullet_point}")

        # Add research experience
        doc.add_heading("Research Experience", level=1)
        for experience in self.research_experience:
            doc.add_heading(f"{experience.position or ''} at {experience.organization or ''}", level=2)
            doc.add_paragraph(f"{experience.start_date or ''} - {experience.end_date or ''}")
            doc.add_paragraph(f"Summary: {experience.summary or ''}")
            for bullet_point in experience.bullet_points or []:
                doc.add_paragraph(f"• {bullet_point}")

        # Add project experience
        doc.add_heading("Project Experience", level=1)
        for experience in self.project_experience:
            doc.add_heading(f"{experience.name or ''}", level=2)
            doc.add_paragraph(f"{experience.start_date or ''} - {experience.end_date or ''}")
            doc.add_paragraph(f"Summary: {experience.summary or ''}")
            for bullet_point in experience.bullet_points or []:
                doc.add_paragraph(f"• {bullet_point}")

        # Add personal statement
        doc.add_heading("Personal Statement", level=1)
        doc.add_paragraph(self.personal_statement or '')

        # Save the document
        doc.save("resume.docx")
        print("Resume generated successfully!")

def main():
    parser = argparse.ArgumentParser(description="Python Resume Generator")
    parser.add_argument("--input", help="Path to input file")

    args = parser.parse_args()

    ui = UserInterface()

    if args.input:
        # Read input from file
        with open(args.input, "r") as file:
            input_data = file.read()

        # Process user input
        # Here, you can implement logic to parse and validate the input data
        # and populate the necessary variables in the UserInterface class
        # For simplicity, I'm assuming the input data is in JSON format
        data = json.loads(input_data)
        ui.personal_info = PersonalInfo(**data.get("personal_info", {}))
        ui.education = [Education(**edu) for edu in data.get("education", [])]
        ui.work_experience = [WorkExperience(**exp) for exp in data.get("work_experience", [])]
        ui.research_experience = [ResearchExperience(**exp) for exp in data.get("research_experience", [])]
        ui.project_experience = [ProjectExperience(**exp) for exp in data.get("project_experience", [])]
        ui.personal_statement = data.get("personal_statement", "")

    else:
        # Capture user input
        ui.capture_personal_info()
        ui.capture_education()
        ui.capture_work_experience()
        ui.capture_research_experience()
        ui.capture_project_experience()
        ui.capture_personal_statement()

    # Display the resume
    ui.display_resume()

if __name__ == "__main__":
    main()