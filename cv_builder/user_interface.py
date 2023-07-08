from docx import Document
from docx.shared import Pt, Inches
from cv_builder.models import (
    Resume,
    PersonalInfo,
    Education,
    WorkExperience,
    ResearchExperience,
    ProjectExperience,
)
import os


class UserInterface:
    def __init__(self, case_path: str = "cases/example", resume: Resume = None):
        if resume:
            self.resume = resume
        else:
            self.resume = Resume(
                personal_info=None,
                education=None,
                work_experience=None,
                research_experience=None,
                project_experience=None,
                personal_statement=None,
            )
        self.case_path = case_path

    def capture_personal_info(self):
        print("Please provide your personal information:")
        self.resume.personal_info = PersonalInfo(
            name=input("Name: "), email=input("Email: "), phone=input("Phone: ")
        )

    def capture_education(self):
        print("Please provide your education details (press enter to skip):")
        while True:
            degree = input("Degree: ")
            if not degree:
                break

            university = input("University: ")
            graduation_date = input("Graduation Date (MM/YYYY): ")

            self.resume.education.append(
                Education(
                    degree=degree, university=university, graduation_date=graduation_date
                )
            )

    def capture_work_experience(self):
        print("Please provide your work experience (press enter to skip):")
        while True:
            company = input("Company: ")
            if not company:
                break

            position = input("Position: ")
            start_date = input("Start Date (MM/YYYY): ")
            end_date = input("End Date (MM/YYYY): ")
            summary = input("Summary: ")

            bullet_points = []
            while True:
                bullet_point = input("Bullet Point (press enter to skip): ")
                if not bullet_point:
                    break
                bullet_points.append(bullet_point)

            self.resume.work_experience.append(
                WorkExperience(
                    company=company,
                    position=position,
                    start_date=start_date,
                    end_date=end_date,
                    summary=summary,
                    bullet_points=bullet_points,
                )
            )

    def capture_research_experience(self):
        print("Please provide your research experience (press enter to skip):")
        while True:
            organization = input("Organization: ")
            if not organization:
                break

            position = input("Position: ")
            start_date = input("Start Date (MM/YYYY): ")
            end_date = input("End Date (MM/YYYY): ")
            summary = input("Summary: ")

            bullet_points = []
            while True:
                bullet_point = input("Bullet Point (press enter to skip): ")
                if not bullet_point:
                    break
                bullet_points.append(bullet_point)

            self.resume.research_experience.append(
                ResearchExperience(
                    organization=organization,
                    position=position,
                    start_date=start_date,
                    end_date=end_date,
                    summary=summary,
                    bullet_points=bullet_points,
                )
            )

    def capture_project_experience(self):
        print("Please provide your project experience (press enter to skip):")
        while True:
            name = input("Project Name: ")
            if not name:
                break

            start_date = input("Start Date (MM/YYYY): ")
            end_date = input("End Date (MM/YYYY): ")
            summary = input("Summary: ")

            bullet_points = []
            while True:
                bullet_point = input("Bullet Point (press enter to skip): ")
                if not bullet_point:
                    break
                bullet_points.append(bullet_point)

            self.resume.project_experience.append(
                ProjectExperience(
                    name=name,
                    start_date=start_date,
                    end_date=end_date,
                    summary=summary,
                    bullet_points=bullet_points,
                )
            )

    def capture_personal_statement(self):
        print("Please provide your personal statement (press enter to skip):")
        self.resume.personal_statement = input("Personal Statement: ")

    def display_resume(self):
        doc = Document()

        # Set document font size and margin
        doc.styles["Normal"].font.size = Pt(10)
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.25)  # Adjust top margin to make space smaller
            section.bottom_margin = Inches(
                0.25
            )  # Adjust bottom margin to make space smaller

        # Add personal information
        doc.add_heading("Personal Information", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Name: {self.resume.personal_info.name or ''}").bold = True
        p.add_run(" | ")
        p.add_run(f"Email: {self.resume.personal_info.email or ''}").bold = True
        p.add_run(" | ")
        p.add_run(f"Phone: {self.resume.personal_info.phone or ''}").bold = True

        # Add education details
        doc.add_heading("Education", level=1)
        for education in self.resume.education or []:
            doc.add_heading(f"{education.degree or ''}", level=2)
            doc.add_paragraph(f"University: {education.university or ''}")
            doc.add_paragraph(f"Graduation Date: {education.graduation_date or ''}")

        # Add work experience
        doc.add_heading("Work Experience", level=1)
        for experience in self.resume.work_experience or []:
            doc.add_heading(
                f"{experience.position or ''} at {experience.company or ''}", level=2
            )
            doc.add_paragraph(
                f"{experience.start_date or ''} - {experience.end_date or ''}"
            )
            doc.add_paragraph(f"Summary: {experience.summary or ''}")
            for bullet_point in experience.bullet_points or []:
                doc.add_paragraph(f"• {bullet_point}")

        # Add research experience
        doc.add_heading("Research Experience", level=1)
        for experience in self.resume.research_experience or []:
            doc.add_heading(
                f"{experience.position or ''} at {experience.organization or ''}", level=2
            )
            doc.add_paragraph(
                f"{experience.start_date or ''} - {experience.end_date or ''}"
            )
            doc.add_paragraph(f"Summary: {experience.summary or ''}")
            for bullet_point in experience.bullet_points or []:
                doc.add_paragraph(f"• {bullet_point}")

        # Add project experience
        doc.add_heading("Project Experience", level=1)
        for experience in self.resume.project_experience or []:
            doc.add_heading(f"{experience.name or ''}", level=2)
            doc.add_paragraph(
                f"{experience.start_date or ''} - {experience.end_date or ''}"
            )
            doc.add_paragraph(f"Summary: {experience.summary or ''}")
            for bullet_point in experience.bullet_points or []:
                doc.add_paragraph(f"• {bullet_point}")

        # Add personal statement
        doc.add_heading("Personal Statement", level=1)
        doc.add_paragraph(self.resume.personal_statement or "")

        # Save the document
        doc.save(os.path.join(self.case_path, "resume.docx"))
        print("Resume generated successfully!")
